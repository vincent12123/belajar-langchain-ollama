from ollama import Client
from docx import Document

OLLAMA_MODEL = "gpt-oss:20b-cloud"
OLLAMA_BASE_URL = "http://localhost:11434"

client = Client(host=OLLAMA_BASE_URL)

def llm(prompt: str, system: str = "") -> str:
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    resp = client.chat(model=OLLAMA_MODEL, messages=messages)
    return resp["message"]["content"].strip()

def make_outline(topic: str, audience: str, style: str, length: str) -> str:
    system = "Kamu penulis dokumen profesional berbahasa Indonesia. Buat struktur yang jelas dan mudah diikuti."
    prompt = f"""
Buat outline dokumen tentang: {topic}
Audiens: {audience}
Gaya bahasa: {style}
Target panjang: {length}

Keluaran WAJIB:
1) Judul (baris pertama, format: "# ...")
2) Daftar isi (Heading 2: "## ...")
3) Catatan singkat tiap bagian (1-2 kalimat)

Formatkan sebagai markdown.
"""
    return llm(prompt, system)

def write_section(topic: str, outline: str, section_title: str, notes: str, style: str) -> str:
    system = "Kamu penulis dokumen profesional. Tulis ringkas, jelas, dan terstruktur."
    prompt = f"""
Topik dokumen: {topic}
Gaya bahasa: {style}

Outline (konteks):
{outline}

Tulis bagian: "{section_title}"
Catatan bagian: {notes}

Aturan:
- Tulis HANYA isi untuk bagian itu (jangan ulang daftar isi)
- Pakai Bahasa Indonesia
- Jika cocok, gunakan bullet/checklist
"""
    return llm(prompt, system)

def polish(full_markdown: str, style: str) -> str:
    system = "Kamu editor Bahasa Indonesia. Rapikan tanpa mengubah maksud."
    prompt = f"""
Rapikan dokumen markdown berikut:
- Buang repetisi
- Perjelas kalimat
- Konsisten istilah
- Pertahankan struktur heading (#, ##, ###)
Gaya bahasa: {style}

DOKUMEN:
{full_markdown}
"""
    return llm(prompt, system)

def save_docx_from_markdown(md: str, filename: str) -> None:
    doc = Document()

    for line in md.splitlines():
        line = line.rstrip()
        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:2].isdigit() and line[2:3] == ".":
            # simple numbered list like "1. ..."
            doc.add_paragraph(line.strip(), style="List Number")
        else:
            doc.add_paragraph(line.strip())

    doc.save(filename)

def build_document(topic: str, audience: str, style: str, length: str, sections: list[tuple[str, str]]) -> str:
    outline = make_outline(topic, audience, style, length)

    parts = [outline, "\n\n"]
    for title, notes in sections:
        section_text = write_section(topic, outline, title, notes, style)
        parts.append(f"## {title}\n{section_text}\n\n")

    draft = "".join(parts)
    final_md = polish(draft, style)
    return final_md

if __name__ == "__main__":
    topic = "SOP Pembuatan Dokumen Menggunakan AI LLM"
    audience = "Tim internal perusahaan"
    style = "formal"
    length = "2-3 halaman"

    # Flow: tulis per bagian (lebih stabil daripada sekali jadi)
    sections = [
        ("Pendahuluan", "Tujuan SOP, definisi singkat LLM, dan ruang lingkup."),
        ("Persiapan Bahan", "Daftar informasi yang wajib disiapkan sebelum minta AI menulis."),
        ("Prosedur Penulisan", "Langkah: outline -> draft per bagian -> revisi -> final."),
        ("Kontrol Kualitas", "Checklist anti-halusinasi: verifikasi angka, istilah, rujukan."),
        ("Keamanan & Kerahasiaan", "Aturan data sensitif, anonimisasi, dan penyimpanan."),
        ("Lampiran", "Template prompt dan contoh format dokumen.")
    ]

    md = build_document(topic, audience, style, length, sections)

    # Simpan
    out_file = "SOP_AI_LLM.docx"
    save_docx_from_markdown(md, out_file)

    print("Selesai! Output:", out_file)
